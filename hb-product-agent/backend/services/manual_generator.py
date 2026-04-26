"""
服务手册生成器（HTML 模板 + python-docx 混合方案）

结构：
- 固定文字内容来自 templates/manual/*.md（可 diff，易编辑）
- 服务段落由 python-docx 编程构建（从 Service 表读取）
- 不依赖二进制 docx 模板
"""
import json
import os
import re
import logging
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from database import GeneratedScheme, Service, GeneratedManual
from config import TEMPLATES_DIR, OUTPUT_MANUALS_DIR

logger = logging.getLogger("manual_generator")


# ============================================================
# 样式常量
# ============================================================
FONT_FAMILY = "微软雅黑"
TITLE_SIZE = Pt(22)
SUBTITLE_SIZE = Pt(14)
HEADING_SIZE = Pt(14)
BODY_SIZE = Pt(10.5)
SMALL_SIZE = Pt(9)

COLOR_DARK = RGBColor(0x00, 0x33, 0x66)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_GRAY = RGBColor(0x90, 0x93, 0x99)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ACCENT = RGBColor(0x00, 0x66, 0xCC)


class ManualGenerator:
    """基于模板 + 数据库的服务手册生成器"""

    def __init__(self):
        self.template_dir = TEMPLATES_DIR
        self.output_dir = OUTPUT_MANUALS_DIR
        os.makedirs(self.output_dir, exist_ok=True)

    # ========== 主入口 ==========

    def generate_manual(
        self, db: Session, scheme_id: int
    ) -> Tuple[Optional[GeneratedManual], List[str]]:
        """
        根据方案生成服务手册 docx
        返回 (GeneratedManual, missing_services)
        """
        scheme = db.query(GeneratedScheme).filter(
            GeneratedScheme.id == scheme_id
        ).first()
        if not scheme:
            raise ValueError("方案不存在")
        if scheme.status != "confirmed":
            raise ValueError("方案未确认，无法生成服务手册")

        # 解析方案服务列表
        scheme_services = self._parse_scheme_services(scheme)
        if not scheme_services:
            raise ValueError("方案中无服务项，无法生成服务手册")

        # 匹配素材库
        db_services = {s.name: s for s in db.query(Service).all()}
        matched, missing = self._match_services(scheme_services, db_services)

        # 计算版本号
        existing_count = db.query(GeneratedManual).filter(
            GeneratedManual.scheme_id == scheme_id
        ).count()
        version = existing_count + 1

        # 生成文件名
        safe_name = self._safe_filename(scheme.scheme_name or "方案")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_name}_v{version}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)

        # 构建文档
        doc = Document()
        self._setup_styles(doc)
        self._render_cover(doc, scheme)
        self._render_warm_tips(doc)
        self._render_services(doc, scheme_services, matched)
        self._render_appendix(doc)
        self._set_page_margins(doc)

        doc.save(filepath)
        logger.info(f"手册已保存: {filepath}")

        # 写入数据库
        manual = GeneratedManual(
            scheme_id=scheme_id,
            version=version,
            manual_title=f"{scheme.scheme_name or '方案'}服务手册",
            docx_path=filepath,
            status="generated",
        )
        db.add(manual)
        db.commit()
        db.refresh(manual)

        return manual, missing

    # ========== 样式设置 ==========

    def _setup_styles(self, doc: Document):
        """设置文档默认样式"""
        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_FAMILY
        font.size = BODY_SIZE
        font.color.rgb = COLOR_BODY
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

        # 段落间距
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5

    def _set_page_margins(self, doc: Document):
        """设置页边距"""
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    # ========== 封面 ==========

    def _render_cover(self, doc: Document, scheme: GeneratedScheme):
        """渲染封面"""
        # 空行留白
        for _ in range(6):
            doc.add_paragraph()

        # 主标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{scheme.scheme_name or '产品方案'}\n健康服务手册")
        run.font.size = TITLE_SIZE
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK
        run.font.name = FONT_FAMILY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

        # 空行
        doc.add_paragraph()

        # 副标题 / 场景描述
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_text = f"（{scheme.scene or '通用场景'} · {scheme.target_group or '目标人群'}）"
        run = subtitle.add_run(sub_text)
        run.font.size = SUBTITLE_SIZE
        run.font.color.rgb = COLOR_GRAY
        run.font.name = FONT_FAMILY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

        # 生成日期
        doc.add_paragraph()
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(datetime.now().strftime("%Y年%m月%d日"))
        run.font.size = BODY_SIZE
        run.font.color.rgb = COLOR_GRAY
        run.font.name = FONT_FAMILY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

        # 分页
        doc.add_page_break()

    # ========== 温馨提示 ==========

    def _render_warm_tips(self, doc: Document):
        """渲染温馨提示"""
        content = self._load_template("warm_tips.md")
        if not content:
            return

        # 标题
        h = doc.add_paragraph()
        run = h.add_run("温馨提示")
        run.font.size = HEADING_SIZE
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK
        run.font.name = FONT_FAMILY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

        # 内容
        lines = content.strip().split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 跳过 markdown 标题
            if line.startswith("## "):
                continue

            # 粗体
            if line.startswith("**") and line.endswith("**"):
                line = line[2:-2]
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.bold = True
            else:
                p = doc.add_paragraph(line)

            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.font.name = FONT_FAMILY
                run.font.size = BODY_SIZE
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

        # 分隔线
        self._add_separator(doc)
        doc.add_paragraph()

    # ========== 服务段落 ==========

    def _render_services(
        self,
        doc: Document,
        scheme_services: List[Dict],
        matched: Dict[str, Any],
    ):
        """渲染服务列表，结构与原版 Word 手册一致"""
        counter = 0
        for svc in scheme_services:
            svc_name = svc.get("name", "")
            db_svc = matched.get(svc_name)
            counter += 1

            # ── 服务标题 ──
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(f"{counter}.{svc_name}")
            title_run.font.size = HEADING_SIZE
            title_run.font.bold = True
            title_run.font.color.rgb = COLOR_ACCENT
            title_run.font.name = FONT_FAMILY
            title_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

            # 服务描述
            description = svc.get("content") or (db_svc.description if db_svc else "")
            if description:
                p = doc.add_paragraph(description)
                p.style = doc.styles["Normal"]

            # ── 服务时间 ──
            svc_time = svc.get("service_time") or (db_svc.service_time if db_svc else "")
            if svc_time:
                self._add_field_line(doc, "服务时间", svc_time)

            # ── 服务时效 ──
            svc_resp = svc.get("service_response_time") or (db_svc.service_response_time if db_svc else "")
            if svc_resp:
                self._add_field_line(doc, "服务时效", svc_resp)

            # ── 服务频次 ──
            times = svc.get("times") or (db_svc.times if db_svc else "")
            if times:
                self._add_field_line(doc, "服务频次", times)

            # ── 启动条件 ──
            condition = svc.get("condition") or (db_svc.condition if db_svc else "")
            if condition:
                self._add_field_line(doc, "启动条件", condition)

            # ── 服务标准 ──
            standard = svc.get("standard") or (db_svc.service_standard if db_svc else "")
            if standard:
                self._add_field_line(doc, "服务标准", standard)

            # ── 服务网络 ──
            network = svc.get("network") or (db_svc.service_network if db_svc else "")
            if network:
                self._add_field_line(doc, "服务网络", network)

            # ── 服务流程 ──
            process = (db_svc.process if db_svc else "")
            if process:
                self._add_field_line(doc, "服务流程", "")
                steps = process.strip().split("\n")
                for step_idx, step in enumerate(steps, 1):
                    step = step.strip()
                    if step:
                        p = doc.add_paragraph(f"（{step_idx}）{step}")
                        p.style = doc.styles["Normal"]

            # ── 特别说明 ──
            notes = svc.get("special_notes") or (db_svc.special_notes if db_svc else "")
            if notes:
                self._add_field_line(doc, "特别说明", "")
                for note_line in notes.strip().split("\n"):
                    note_line = note_line.strip()
                    if note_line:
                        p = doc.add_paragraph(note_line)
                        p.style = doc.styles["Normal"]

            # ── 价格参考 ──
            cost = svc.get("cost_price") or svc.get("cost") or ""
            quote = svc.get("quote_price") or svc.get("price") or svc.get("quote") or ""
            if cost or quote:
                parts = []
                if cost:
                    parts.append(f"成本 {cost}元")
                if quote:
                    parts.append(f"报价 {quote}元")
                self._add_field_line(doc, "参考价格", " / ".join(parts))

            doc.add_paragraph()  # 服务间空行

    def _add_field_line(self, doc: Document, label: str, value: str):
        """添加「标签：值」格式的段落"""
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}：")
        run_label.font.size = BODY_SIZE
        run_label.font.bold = True
        run_label.font.name = FONT_FAMILY
        run_label._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
        if value:
            run_value = p.add_run(value)
            run_value.font.size = BODY_SIZE
            run_value.font.name = FONT_FAMILY
            run_value._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)

    # ========== 附件 ==========

    def _render_appendix(self, doc: Document):
        """渲染附件"""
        self._add_separator(doc)

        content = self._load_template("appendix.md")
        if not content:
            return

        lines = content.strip().split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith("## "):
                h = doc.add_paragraph()
                run = h.add_run(line[3:])
                run.font.size = HEADING_SIZE
                run.font.bold = True
                run.font.color.rgb = COLOR_DARK
                run.font.name = FONT_FAMILY
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
            elif line.startswith("### "):
                h = doc.add_paragraph()
                run = h.add_run(line[4:])
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = COLOR_DARK
                run.font.name = FONT_FAMILY
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
            elif re.match(r"^\d+\.\s", line):
                p = doc.add_paragraph(line)
                p.style = doc.styles["Normal"]
                p.paragraph_format.left_indent = Cm(0.5)
            else:
                p = doc.add_paragraph(line)
                p.style = doc.styles["Normal"]

    # ========== 工具方法 ==========

    def _load_template(self, filename: str) -> Optional[str]:
        """读取模板文件内容"""
        path = os.path.join(self.template_dir, filename)
        if not os.path.exists(path):
            logger.warning(f"模板文件不存在: {path}")
            return None
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    def _parse_scheme_services(self, scheme: GeneratedScheme) -> List[Dict[str, Any]]:
        """从 GeneratedScheme 解析服务列表，合并所有方案的服务"""
        services = []
        if not scheme.service_list_json:
            return services

        try:
            parsed = json.loads(scheme.service_list_json)
        except (json.JSONDecodeError, TypeError):
            return services

        if isinstance(parsed, dict):
            # 多方案：从 schemes 数组中收集所有服务
            schemes = parsed.get("schemes", [])
            if schemes:
                seen = set()
                for sch in schemes:
                    if not isinstance(sch, dict):
                        continue
                    sch_name = sch.get("scheme_name", "")
                    for svc in sch.get("services", []) or []:
                        if not isinstance(svc, dict):
                            logger.warning(f"[_parse_scheme_services] 跳过非dict服务项: {str(svc)[:80]}")
                            continue
                        name = svc.get("name", "")
                        if name and name not in seen:
                            seen.add(name)
                            svc_copy = dict(svc)
                            if sch_name:
                                svc_copy["_scheme_group"] = sch_name
                            services.append(svc_copy)
            # 兼容单方案：顶层 services
            if not services:
                services = [s for s in (parsed.get("services", []) or []) if isinstance(s, dict)]
        elif isinstance(parsed, list):
            services = [s for s in parsed if isinstance(s, dict)]

        return services

    def _match_services(
        self,
        scheme_services: List[Dict[str, Any]],
        db_services: Dict[str, Service],
    ) -> Tuple[Dict[str, Service], List[str]]:
        """
        将方案服务与素材库匹配
        返回 ({方案服务名 → Service对象}, 未匹配名称列表)
        """
        matched = {}
        missing = []

        for svc in scheme_services:
            svc_name = svc.get("name", "")
            if not svc_name:
                continue

            db_svc = self._find_db_service(svc_name, db_services, set(matched.values()))
            if db_svc:
                matched[svc_name] = db_svc
            else:
                missing.append(svc_name)

        return matched, missing

    def _find_db_service(
        self,
        svc_name: str,
        db_services: Dict[str, Service],
        used: set,
    ) -> Optional[Service]:
        """在素材库中查找匹配的服务"""
        # 1. 精确匹配
        if svc_name in db_services and db_services[svc_name] not in used:
            return db_services[svc_name]

        # 2. 子串匹配
        for db_name, db_svc in db_services.items():
            if db_svc in used:
                continue
            if svc_name in db_name or db_name in svc_name:
                return db_svc

        # 3. 别名匹配
        for db_name, db_svc in db_services.items():
            if db_svc in used:
                continue
            aliases = (db_svc.aliases or "").split(",")
            for alias in aliases:
                alias = alias.strip()
                if alias and (alias in svc_name or svc_name in alias):
                    return db_svc

        # 4. 核心词匹配（去常见后缀）
        svc_core = svc_name.replace("服务", "").replace("权益", "")
        for db_name, db_svc in db_services.items():
            if db_svc in used:
                continue
            db_core = db_name.replace("服务", "").replace("权益", "")
            if len(svc_core) >= 2 and len(db_core) >= 2:
                if svc_core in db_core or db_core in svc_core:
                    return db_svc

        return None

    def _safe_filename(self, name: str) -> str:
        """将方案名转为安全文件名"""
        safe = re.sub(r"[^\w\u4e00-\u9fff\-]", "_", name)
        safe = re.sub(r"_+", "_", safe).strip("_")
        return safe[:60] if safe else "scheme"

    def _add_separator(self, doc: Document):
        """添加分隔线"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("—" * 30)
        run.font.size = SMALL_SIZE
        run.font.color.rgb = COLOR_GRAY
